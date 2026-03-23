import io
from typing import Dict

import streamlit as st
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

st.set_page_config(page_title="PAGE BUILDER V2", layout="wide")

if "reset_nonce" not in st.session_state:
    st.session_state.reset_nonce = 0

st.title("PAGE BUILDER V2")
st.caption("AI 검색 최적화형 상세페이지 + SEO 콘텐츠 생성기")

api_key = st.secrets.get("OPENAI_API_KEY", "")
if not api_key:
    st.warning("OPENAI_API_KEY가 설정되지 않았습니다.")
    st.stop()

client = OpenAI(api_key=api_key)

PROMPT = """
너는 4050 여성 패션 전문 온라인몰 ‘미샵(MISHARP)’의 시니어 에디터이자 SEO 콘텐츠 전략가다.

목표:
1. 검색엔진과 AI 검색(구글 AI Overviews, 네이버 AI 브리핑)에 잘 이해되는 구조로
2. 실제 구매전환이 일어날 수 있도록
3. 미샵의 고급스럽고 신뢰감 있는 브랜드 톤으로
상품 상세페이지 원고를 작성한다.

핵심 원칙:
- 단순 상품 소개가 아니라 “고객 문제를 해결하는 정보형 상품 콘텐츠”로 작성한다.
- 키워드 반복보다 사용자에게 유용한 설명을 우선한다.
- 문장은 짧고 명확하게 작성한다.
- 과장, 허위, 검증 불가 표현은 피한다.
- 검색엔진이 이해하기 쉽도록 구조화된 소제목 체계를 사용한다.
- 4050 여성의 체형 고민, 생활 동선, 스타일 고민을 실제적으로 반영한다.
- 미샵 특유의 감성은 살리되, 핵심 정보 전달이 흐려지지 않도록 한다.
- 작성 시 가장 중요한 기준은 “사용자가 검색창에 입력할 질문에 실제로 답이 되는가”이다.
- 상품명 반복보다 검색의도 해결을 우선하라.
- 브랜드 홍보문처럼만 쓰지 말고, 실제 구매 전 판단에 도움이 되는 정보로 써라.

출력 형식:
아래 순서를 반드시 지켜 작성한다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SEO 제목 3개
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
- 검색형 제목
- 문제해결형 제목
- 감성+검색 혼합형 제목

조건:
- 각 제목은 28~38자 내외
- 핵심 키워드가 자연스럽게 포함되게 작성
- 상품명만 반복하지 말 것

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
메타디스크립션 2개
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
- 각 70~110자 내외
- 클릭하고 싶게 작성
- 고객 문제와 활용 상황이 드러나야 함

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
AI 검색 친화 요약문
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
- 3~5문장
- 검색 결과 상단 요약에 인용되어도 자연스러운 방식
- 이 제품이 어떤 문제를 해결하고 어떤 상황에서 좋은지가 분명해야 함

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
상세페이지 본문 구조
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
H2. 이 상품을 추천하는 이유
H2. 원단과 소재의 장점
H2. 체형과 핏에 대하여
H2. 이런 분들께 추천합니다
H2. 이렇게 코디해보세요
H2. 자주 묻는 질문
H2. 리뷰 요약

각 H2 아래 내용은 미샵 톤으로 자연스럽게 작성한다.
FAQ는 실제 검색 질문처럼 4개 작성한다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
이미지 ALT 텍스트 6개
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
내부링크용 앵커텍스트 5개
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
구조화데이터 초안용 정보 정리
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
- product name
- description
- brand
- color
- material
- size
- category
- url

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
최종 점검 체크리스트
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
- 검색의도 반영 여부
- 체형 고민 반영 여부
- 코디 활용도 반영 여부
- FAQ 포함 여부
- AI 검색용 요약 포함 여부
- ALT 텍스트 포함 여부

중요:
- 일반 마크다운 코드펜스 금지
- 실무 문서처럼 바로 복사 가능한 텍스트로 작성
"""

def build_prompt(data: Dict[str, str]) -> str:
    return f"""{PROMPT}

입력 정보:
- 상품명: {data['product_name']}
- 카테고리: {data['category']}
- 소재: {data['material']}
- 컬러: {data['color']}
- 사이즈: {data['size']}
- 핏 특징: {data['fit']}
- 디테일 특징: {data['detail']}
- 추천 코디: {data['coordi']}
- 추천 고객: {data['target']}
- 착용 상황/TPO: {data['tpo']}
- 가격/할인 정보: {data['price_info']}
- 리뷰/고객 반응: {data['review_summary']}
- 상품 URL: {data['url']}
- 핵심 검색 질문 3~5개: {data['search_questions']}
- 고객 문제: {data['customer_problem']}
- 기타 참고사항: {data['etc']}
"""

def result_to_docx_bytes(result_text: str) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    style.font.size = Pt(10.5)

    for line in result_text.splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Malgun Gothic"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        run.font.size = Pt(10.5)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

h1, h2, h3 = st.columns([2.4, 1.0, 8.6])
with h1:
    st.subheader("정보 입력")
with h2:
    st.write("")
    if st.button("초기화", use_container_width=True):
        st.session_state.reset_nonce += 1
        st.rerun()

nonce = st.session_state.reset_nonce
left, right = st.columns(2)

with left:
    product_name = st.text_input("상품명", placeholder="예: 볼륨 셔링 에코 레더 조끼", key=f"product_name_{nonce}")
    category = st.text_input("카테고리", placeholder="예: 베스트 / 자켓 / 슬랙스", key=f"category_{nonce}")
    material = st.text_input("소재", placeholder="예: 에코 레더 / 면80 나일론20", key=f"material_{nonce}")
    color = st.text_input("컬러", placeholder="예: 블랙, 브라운", key=f"color_{nonce}")
    size = st.text_input("사이즈", placeholder="예: Free / S~XL", key=f"size_{nonce}")
    fit = st.text_input("핏 특징", placeholder="예: 복부 커버되는 여유핏", key=f"fit_{nonce}")
    detail = st.text_area("디테일 특징", height=110, placeholder="예: 셔링 밑단, 절개라인, 인밴딩", key=f"detail_{nonce}")

with right:
    search_questions = st.text_area("핵심 검색 질문 3~5개", height=120, placeholder="예: 가죽조끼 부해 보이지 않나요?\n40대 레더조끼 코디 어떻게 하나요?", key=f"search_questions_{nonce}")
    customer_problem = st.text_area("고객 문제", height=110, placeholder="예: 복부 커버, 팔뚝 커버, 출근룩 고민", key=f"customer_problem_{nonce}")
    target = st.text_input("추천 고객", value="4050 여성", key=f"target_{nonce}")
    tpo = st.text_input("착용 상황/TPO", placeholder="예: 출근룩, 모임룩, 학교방문룩", key=f"tpo_{nonce}")
    coordi = st.text_area("추천 코디", height=90, placeholder="예: 화이트 티 + 데님 / 셔츠 + 슬랙스", key=f"coordi_{nonce}")
    price_info = st.text_input("가격/할인 정보", placeholder="예: 신상 10% 할인", key=f"price_info_{nonce}")
    review_summary = st.text_area("리뷰/고객 반응", height=90, placeholder='예: "부드럽고 부담 없어요" / "배가 덜 드러나요"', key=f"review_summary_{nonce}")
    url = st.text_input("상품 URL", placeholder="예: https://www.misharp.co.kr/...", key=f"url_{nonce}")
    etc = st.text_area("기타 참고사항", height=90, placeholder="예: 브랜드 톤은 차분하고 세련되게", key=f"etc_{nonce}")

if st.button("V2 생성하기", type="primary", use_container_width=True, key=f"generate_{nonce}"):
    if not product_name.strip():
        st.warning("상품명을 입력해 주세요.")
        st.stop()

    data = {
        "product_name": product_name,
        "category": category,
        "material": material,
        "color": color,
        "size": size,
        "fit": fit,
        "detail": detail,
        "coordi": coordi,
        "target": target,
        "tpo": tpo,
        "price_info": price_info,
        "review_summary": review_summary,
        "url": url,
        "search_questions": search_questions,
        "customer_problem": customer_problem,
        "etc": etc,
    }

    prompt = build_prompt(data)

    with st.spinner("V2 출력물을 생성 중입니다..."):
        response = client.chat.completions.create(
            model="gpt-4.1",
            messages=[
                {"role": "system", "content": "너는 미샵의 AI 검색 최적화형 SEO/상세페이지 전략가다. 구조와 순서를 반드시 지켜라."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.25,
        )
        result = response.choices[0].message.content

    st.success("생성이 완료되었습니다.")
    st.text_area("V2 결과", result, height=1100, key=f"result_{nonce}")

    docx_bytes = result_to_docx_bytes(result)
    d1, d2 = st.columns(2)
    with d1:
        st.download_button(
            "TXT 다운로드",
            data=result,
            file_name=f"{product_name}_page_builder_v2.txt",
            mime="text/plain",
            use_container_width=True,
            key=f"download_txt_{nonce}"
        )
    with d2:
        st.download_button(
            "DOCX 다운로드",
            data=docx_bytes,
            file_name=f"{product_name}_page_builder_v2.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key=f"download_docx_{nonce}"
        )

st.markdown("---")
st.markdown("© made by MISHARP, MIYAWA")
